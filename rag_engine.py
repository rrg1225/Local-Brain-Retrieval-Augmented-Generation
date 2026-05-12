"""
RAG 核心：本地 HuggingFace GPU Embedding + SentenceTransformer Rerank、
BM25+Vector 混合检索 (RRF)、Chroma 多 Collection 物理隔离、
基于 mtime 的增量索引缓存、分批防抖落盘，
以及基于 Condense + Context 的 Chat Engine。

扩展：Java / Vue 代码感知切块、图片 Gemini 视觉描述、工作区忽略目录剪枝、
并发建库（主线程批量写入 Chroma）、PyMuPDF 解析 PDF / python-docx 解析 Word。
"""

from __future__ import annotations

import json
import mimetypes
import os
import re
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Callable, Iterable

import chromadb
import docx
import fitz
import jieba
import torch
from google.genai import types as genai_types

from llama_index.core import Document, Settings, StorageContext, VectorStoreIndex
from llama_index.core.chat_engine.condense_plus_context import CondensePlusContextChatEngine
from llama_index.core.chat_engine.types import BaseChatEngine, ChatMode
from llama_index.core.llms import LLM
from llama_index.core.memory import ChatMemoryBuffer
from llama_index.core.node_parser import CodeSplitter
from llama_index.core.node_parser.node_utils import build_nodes_from_splits, default_id_func
from llama_index.core.node_parser.text.sentence import SentenceSplitter
from llama_index.core.prompts import PromptTemplate
from llama_index.core.retrievers import QueryFusionRetriever, VectorIndexRetriever
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.google_genai import GoogleGenAI
from llama_index.postprocessor.sbert_rerank import SentenceTransformerRerank
from llama_index.retrievers.bm25 import BM25Retriever
from llama_index.vector_stores.chroma import ChromaVectorStore

from config import AppConfig, get_config

# ==========================================
# 🚀 终极优化 1：网络与环境配置 (防掉线)
# ==========================================
# 禁用 HTTP/2，彻底解决使用代理时的 [SSL: UNEXPECTED_EOF_WHILE_READING] 报错
os.environ.setdefault("HTTPX_NO_HTTP2", "1")
# 抑制 LlamaIndex 遥测日志
os.environ.setdefault("ANONYMIZED_TELEMETRY", "False")

# ==========================================
# 🚀 终极优化 2：代码级精准分词器
# ==========================================
def code_aware_tokenize(text: str) -> list[str]:
    """
    专为代码库设计的混合分词器：
    1. 完整保留 camelCase, snake_case 等代码变量名。
    2. 提取纯中文注释和说明。
    """
    code_tokens = re.findall(r"[a-zA-Z0-9_]+", text)
    chinese_tokens = [w for w in jieba.lcut(text) if re.search(r"[一-龥]", w)]
    return code_tokens + chinese_tokens

# ==========================================
# 🚀 终极优化 3：BM25 检索引擎全局缓存
# ==========================================
# 格式: { "空间名": (状态文件修改时间戳, BM25Retriever实例) }
_BM25_CACHE: dict[str, tuple[float, BM25Retriever]] = {}

# --- 索引白名单 ---
TEXT_EXTENSIONS = {".md", ".pdf", ".py", ".ts", ".tsx", ".json", ".yaml", ".yml", ".docx"}
CODE_EXTENSIONS = {".java", ".vue"}
IMAGE_EXTENSIONS = {".jpg", ".png", ".jpeg"}
INDEX_EXTENSIONS = TEXT_EXTENSIONS | CODE_EXTENSIONS | IMAGE_EXTENSIONS

# --- 强制忽略的目录名（大小写不敏感）---
IGNORE_DIR_NAMES_LOWER = frozenset(
    name.lower() for name in ("node_modules", "target", ".git", ".idea")
)

# --- 无后缀名但需被索引的文件名（如 Dockerfile、Makefile）---
NO_EXTENSION_FILENAMES = frozenset({
    "Dockerfile", "Makefile", "LICENSE", ".env", ".gitignore",
    ".dockerignore", ".editorconfig", "Vagrantfile", "Gemfile",
    "Rakefile", "Procfile", "Jenkinsfile",
})

# 向量粗检条数；经 Rerank 后保留 cfg.top_k 条
RERANK_COARSE_TOP_K = 20

# 分批写入：每批文件数（控制 GPU 显存峰值）
INITIAL_SCAN_BATCH = 5

EXPERT_SYSTEM_PROMPT = (
    "你是「第二大脑」助手，同时也是一位资深系统架构师和代码审查专家。\n\n"
    "**核心原则：**\n"
    "1. 必须严格基于检索到的本地上下文作答。如果上下文中没有相关信息，请直接说明"
    "「本地代码未包含此信息」，严禁猜测或编造。\n"
    "2. 回答时明确引用涉及的本地文件名和关键符号"
    "（例如 UserService.java、App.vue、getUserById 方法），便于开发者快速定位源码。\n"
    "3. 使用清晰的 Markdown 格式组织回答，代码片段使用 ``` 代码块包裹并标注语言。\n"
    "4. 分析架构或代码逻辑时，指出模块之间的依赖关系和调用链。\n"
    "5. 审查代码时，关注潜在的并发安全、空指针、资源泄露等工程风险。"
)

CAPTION_SYSTEM_INSTRUCTION = (
    "你是资深软件工程师，正在索引项目里的截图与架构图。"
    "请用中文客观描述图片内容：可见的文字、模块名、箭头/连线、数据库表、框架或流程；"
    "若与软件架构相关请点明。不要臆测看不清的细节。"
)

_genai_client = None


# ============================================================================
# 文件遍历与过滤
# ============================================================================

def path_should_skip(path: Path) -> bool:
    """路径任一段落在忽略目录内则跳过（含文件与其父目录链）。"""
    return any(part.lower() in IGNORE_DIR_NAMES_LOWER for part in path.parts)


def walk_workspace_files(workspaces: list[Path]):
    """os.walk + 目录剪枝，避免深入 node_modules / target 等。"""
    for root in workspaces:
        if not root.exists():
            continue
        root = root.resolve()
        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [
                d
                for d in dirnames
                if d.lower() not in IGNORE_DIR_NAMES_LOWER
            ]
            for fn in filenames:
                yield Path(dirpath) / fn


def build_http_options(proxy_base: str) -> genai_types.HttpOptions | None:
    if not proxy_base:
        return None
    return genai_types.HttpOptions(base_url=proxy_base)


def _get_genai_client(cfg: AppConfig):
    """原生 google.genai Client：与 LlamaIndex 共用 base_url / 密钥，便于多模态 caption。"""
    global _genai_client
    if _genai_client is None:
        import google.genai as genai

        _genai_client = genai.Client(
            api_key=cfg.api_key or None,
            http_options=build_http_options(cfg.proxy_base),
        )
    return _genai_client


# ============================================================================
# 切块工具
# ============================================================================

def _code_max_chars(cfg: AppConfig) -> int:
    """将 token 级 chunk_size 粗略映射到 CodeSplitter 的字符上限。"""
    return max(1200, min(6000, cfg.chunk_size * 8))


def _merge_under_limit(parts: list[str], max_chars: int) -> list[str]:
    """合并过碎片段，避免大量超小 node。"""
    merged: list[str] = []
    buf = ""
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if not buf:
            buf = p
            continue
        if len(buf) + 2 + len(p) <= max_chars:
            buf = f"{buf}\n\n{p}"
        else:
            merged.append(buf)
            buf = p
    if buf:
        merged.append(buf)
    return merged


def _hard_wrap(text: str, max_chars: int, overlap: int) -> list[str]:
    """超长块按行硬切，带简单重叠。"""
    if len(text) <= max_chars:
        return [text] if text.strip() else []
    lines = text.split("\n")
    out: list[str] = []
    buf = ""
    for line in lines:
        if len(buf) + len(line) + 1 <= max_chars:
            buf = f"{buf}\n{line}" if buf else line
        else:
            if buf.strip():
                out.append(buf.strip())
            buf = line
    if buf.strip():
        out.append(buf.strip())
    if overlap <= 0 or len(out) <= 1:
        return out
    joined: list[str] = []
    tail_lines = max(1, overlap // 40)
    prev_tail = ""
    for i, chunk in enumerate(out):
        if i == 0:
            joined.append(chunk)
            prev_tail = "\n".join(chunk.split("\n")[-tail_lines:])
        else:
            merged = f"{prev_tail}\n{chunk}" if prev_tail else chunk
            joined.append(merged.strip())
            prev_tail = "\n".join(merged.split("\n")[-tail_lines:])
    return joined


def split_java_regex(text: str, max_chars: int, overlap: int) -> list[str]:
    """Java 回退：在类/接口/方法/注解等边界处切开，减少切碎类名与方法签名。"""
    pattern = re.compile(
        r"(?=\n(?:(?:@\w+(?:\([^)]*\))?\s*)*\s*)*"
        r"(?:public|private|protected|static|final|abstract|synchronized|native|strictfp|class|interface|enum)\b)",
        re.MULTILINE,
    )
    parts = pattern.split(text)
    parts = [p.strip() for p in parts if p.strip()]
    merged = _merge_under_limit(parts, max_chars)
    final: list[str] = []
    for m in merged:
        final.extend(_hard_wrap(m, max_chars, overlap))
    return [x for x in final if x.strip()]


def split_vue_sfc_regex(text: str, max_chars: int, overlap: int) -> list[str]:
    """Vue SFC 回退：按 template/script/style 与常见导出边界切分。"""
    block_pat = re.compile(r"(?=\n?<(?:template|script|style)\b)", re.IGNORECASE)
    blocks = block_pat.split(text)
    inner_pat = re.compile(
        r"(?=\n(?:export\s+default|function\s+\w|const\s+\w+\s*=\s*(?:async\s*)?\(|"
        r"methods\s*:\s*\{|computed\s*:\s*\{|watch\s*:\s*\{))",
    )
    pieces: list[str] = []
    for b in blocks:
        b = b.strip()
        if not b:
            continue
        subs = inner_pat.split(b)
        for s in subs:
            s = s.strip()
            if s:
                pieces.append(s)
    merged = _merge_under_limit(pieces, max_chars)
    final: list[str] = []
    for m in merged:
        final.extend(_hard_wrap(m, max_chars, overlap))
    return [x for x in final if x.strip()]


def _nodes_from_text_splits(doc: Document, splits: list[str]) -> list:
    if not splits:
        return []
    return build_nodes_from_splits(splits, doc, id_func=default_id_func)


def get_code_aware_nodes(doc: Document, path: Path, cfg: AppConfig) -> list:
    """.java / .vue 优先 CodeSplitter（tree-sitter），失败则正则结构切分。"""
    suffix = path.suffix.lower()
    max_c = _code_max_chars(cfg)
    overlap = cfg.chunk_overlap

    if suffix == ".java":
        try:
            splitter = CodeSplitter.from_defaults(
                language="java",
                max_chars=max_c,
                chunk_lines=64,
                chunk_lines_overlap=min(24, max(8, cfg.chunk_overlap // 4)),
            )
            return splitter.get_nodes_from_documents([doc])
        except Exception:
            return _nodes_from_text_splits(doc, split_java_regex(doc.text, max_c, overlap))

    if suffix == ".vue":
        try:
            splitter = CodeSplitter.from_defaults(
                language="vue",
                max_chars=max_c,
                chunk_lines=64,
                chunk_lines_overlap=min(24, max(8, cfg.chunk_overlap // 4)),
            )
            return splitter.get_nodes_from_documents([doc])
        except Exception:
            return _nodes_from_text_splits(doc, split_vue_sfc_regex(doc.text, max_c, overlap))

    splitter = SentenceSplitter(chunk_size=cfg.chunk_size, chunk_overlap=cfg.chunk_overlap)
    return splitter.get_nodes_from_documents([doc])


# ============================================================================
# 文档读取
# ============================================================================

def caption_image_with_gemini(path: Path, cfg: AppConfig) -> str | None:
    try:
        data = path.read_bytes()
    except OSError:
        return None
    mime, _ = mimetypes.guess_type(str(path))
    if not mime:
        mime = "image/jpeg" if path.suffix.lower() in (".jpg", ".jpeg") else "image/png"

    client = _get_genai_client(cfg)
    try:
        image_part = genai_types.Part.from_bytes(data=data, mime_type=mime)
        response = client.models.generate_content(
            model=cfg.llm_model,
            contents=[
                CAPTION_SYSTEM_INSTRUCTION + "\n请输出一段可用于检索的纯文本描述（勿用 Markdown 标题）。",
                image_part,
            ],
        )
        text = (response.text or "").strip()
        return text or None
    except Exception:
        return None


def _read_pdf_pymupdf(path: Path) -> str | None:
    """使用 PyMuPDF (fitz) 提取 PDF 文本，保留顺序与复杂排版能力更强。"""
    try:
        doc = fitz.open(str(path))
        try:
            parts: list[str] = []
            for page in doc:
                parts.append(page.get_text("text", sort=True) or "")
            return "\n".join(parts)
        finally:
            doc.close()
    except (RuntimeError, ValueError, OSError):
        return None


def _read_docx(path: Path) -> str | None:
    try:
        doc = docx.Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def read_text_document(path: Path) -> Document | None:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _read_pdf_pymupdf(path)
        if text is None:
            return None
    elif suffix == ".docx":
        text = _read_docx(path)
        if text is None:
            return None
    else:
        try:
            text = path.read_text(encoding="utf-8", errors="ignore")
        except (OSError, UnicodeError):
            return None
    if not text.strip():
        return None
    doc_id = str(path.resolve())
    return Document(
        text=text,
        metadata={"file_path": doc_id, "file_name": path.name, "content_type": "text"},
        id_=doc_id,
    )


def read_image_caption_document(path: Path, cfg: AppConfig) -> Document | None:
    caption = caption_image_with_gemini(path, cfg)
    if not caption:
        return None
    doc_id = str(path.resolve())
    text = (
        f"[图片文件: {path.name}]\n"
        f"[类型: 由 Gemini 视觉模型生成的索引描述，用于检索]\n\n"
        f"{caption}"
    )
    return Document(
        text=text,
        metadata={
            "file_path": doc_id,
            "file_name": path.name,
            "content_type": "image_caption",
        },
        id_=doc_id,
    )


def build_index_document(path: Path, cfg: AppConfig) -> Document | None:
    suffix = path.suffix.lower()
    if not path.is_file():
        return None
    if suffix not in INDEX_EXTENSIONS and path.name not in NO_EXTENSION_FILENAMES:
        return None
    if suffix in IMAGE_EXTENSIONS:
        return read_image_caption_document(path, cfg)
    return read_text_document(path)


def _nodes_from_built_document(doc: Document, file_path: Path, cfg: AppConfig) -> list:
    """与 incremental_upsert_file 一致的分块策略（供并发 worker 使用）。"""
    suffix = file_path.suffix.lower()
    if suffix in CODE_EXTENSIONS:
        return get_code_aware_nodes(doc, file_path, cfg)
    splitter = SentenceSplitter(chunk_size=cfg.chunk_size, chunk_overlap=cfg.chunk_overlap)
    return splitter.get_nodes_from_documents([doc])


# ============================================================================
# 并发解析 Worker（纯本地切块）
# ============================================================================

def _ingest_worker_task(file_path: Path, cfg: AppConfig) -> tuple[str, list] | None:
    """读文件、切块，构建 IndexNode；Embedding 在 insert_nodes 阶段由本地 GPU 完成。"""
    if path_should_skip(file_path):
        return None
    doc = build_index_document(file_path, cfg)
    if doc is None:
        return None
    doc_id = str(file_path.resolve())
    nodes = _nodes_from_built_document(doc, file_path, cfg)
    if not nodes:
        return None
    return doc_id, nodes


def _initial_scan_max_workers(num_files: int) -> int:
    cpu = os.cpu_count() or 4
    return max(1, min(32, num_files, cpu * 2))


# ============================================================================
# 安全写入（本地 GPU，保留通用异常处理）
# ============================================================================

def safe_insert_nodes(index: VectorStoreIndex, nodes: list, label: str = "") -> bool:
    """写入索引。成功返回 True，失败返回 False。"""
    try:
        index.insert_nodes(nodes)
        return True
    except Exception as e:
        print(f"[第二大脑] {label} 写入异常：{e}")
        return False


# ============================================================================
# 索引状态持久化（按空间隔离：.index_state_{space}.json）
# ============================================================================

def _index_state_path(collection_name: str) -> Path:
    return Path(f".index_state_{collection_name}.json")


def _load_index_state(collection_name: str) -> dict:
    """加载索引状态文件，容错文件不存在或 JSON 损坏。"""
    path = _index_state_path(collection_name)
    try:
        if path.exists():
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                return data
    except (json.JSONDecodeError, OSError):
        pass
    return {}


def _save_index_state(state: dict, collection_name: str) -> None:
    """原子写入：先写临时文件再 rename。"""
    path = _index_state_path(collection_name)
    tmp = path.with_suffix(".tmp")
    try:
        tmp.write_text(json.dumps(state, ensure_ascii=False, indent=2), encoding="utf-8")
        tmp.replace(path)
    except OSError:
        pass


def _prune_stale_entries(state: dict, current_keys: set[str]) -> dict:
    """移除已不存在于磁盘的文件记录。"""
    return {k: v for k, v in state.items() if k in current_keys}


def _update_state_entry(state: dict, key: str, mtime: float, ok: bool) -> None:
    """就地更新单条状态记录。"""
    entry = state.setdefault(key, {})
    entry["mtime"] = mtime
    if ok:
        entry["in_index"] = True


# ============================================================================
# 本地 GPU Embedding 与 Index 构建
# ============================================================================

def _make_local_embed(cfg: AppConfig) -> HuggingFaceEmbedding:
    device = "cuda:0" if torch.cuda.is_available() else "cpu"
    print(f"[第二大脑] 加载本地 Embedding 模型 {cfg.local_embed_model} → {device}")
    try:
        return HuggingFaceEmbedding(
            model_name=cfg.local_embed_model,
            device=device,
            trust_remote_code=True,
        )
    except Exception as e:
        if "onnx" in str(e).lower():
            print(f"[第二大脑] ONNX 推理后端不可用，回退 PyTorch 原生推理：{e}")
            os.environ.setdefault("ORT_DISABLE", "1")
        raise


def _make_gemini_llm(cfg: AppConfig, model_name: str | None = None) -> GoogleGenAI:
    http_opts = build_http_options(cfg.proxy_base)
    return GoogleGenAI(
        model=model_name or cfg.llm_model,
        api_key=cfg.api_key or None,
        http_options=http_opts,
    )


def _get_or_create_index(cfg: AppConfig, embed_model: HuggingFaceEmbedding,
                         collection_name: str = "default_space") -> VectorStoreIndex:
    Path(cfg.chroma_dir).mkdir(parents=True, exist_ok=True)
    db = chromadb.PersistentClient(path=cfg.chroma_dir)
    collection = db.get_or_create_collection(collection_name)
    vector_store = ChromaVectorStore(chroma_collection=collection)
    storage_context = StorageContext.from_defaults(vector_store=vector_store)
    return VectorStoreIndex.from_vector_store(
        vector_store,
        storage_context=storage_context,
        embed_model=embed_model,
    )


def list_existing_collections(cfg: AppConfig) -> list[str]:
    """列出 ChromaDB 中已有的所有 Collection 名称。"""
    try:
        Path(cfg.chroma_dir).mkdir(parents=True, exist_ok=True)
        db = chromadb.PersistentClient(path=cfg.chroma_dir)
        return [c.name for c in db.list_collections()]
    except Exception:
        return ["default_space"]


def delete_workspace_nodes(index: VectorStoreIndex, collection_name: str,
                          workspace_path: Path) -> int:
    """从 index_state 中定位 workspace_path 下所有已索引文件，
    逐一删除向量节点并从状态文件中移除记录。返回清理的文件数。"""
    state = _load_index_state(collection_name)
    resolved = str(workspace_path.resolve())
    keys_to_remove = [k for k in state if k.startswith(resolved)]

    count = 0
    for doc_id in keys_to_remove:
        try:
            index.delete_ref_doc(doc_id, delete_from_docstore=True)
            count += 1
        except Exception:
            pass
        state.pop(doc_id, None)

    _save_index_state(state, collection_name)
    return count


def delete_project_space(cfg: AppConfig, collection_name: str) -> None:
    """删除 ChromaDB 中的 Collection 及其索引状态文件。"""
    try:
        db = chromadb.PersistentClient(path=cfg.chroma_dir)
        db.delete_collection(collection_name)
        print(f"[第二大脑] 已删除 Collection：{collection_name}")
    except Exception as e:
        print(f"[第二大脑] 删除 Collection 失败：{e}")
    state_file = _index_state_path(collection_name)
    try:
        if state_file.exists():
            state_file.unlink()
    except OSError:
        pass


def cleanup_ghost_nodes(index: VectorStoreIndex, collection_name: str) -> int:
    """扫描 index_state 中所有记录，删除已被外部移除的幽灵节点。返回清理数。"""
    state = _load_index_state(collection_name)
    ghost_keys = [k for k in state if not os.path.exists(k)]

    count = 0
    for doc_id in ghost_keys:
        try:
            index.delete_ref_doc(doc_id, delete_from_docstore=True)
            count += 1
        except Exception:
            pass
        state.pop(doc_id, None)

    if ghost_keys:
        _save_index_state(state, collection_name)
        print(f"[第二大脑] 清理 {count} 个幽灵节点")

    if torch.cuda.is_available():
        torch.cuda.empty_cache()

    return count


# ============================================================================
# 单轨入库
# ============================================================================

def incremental_upsert_file(index: VectorStoreIndex, cfg: AppConfig, file_path: Path,
                            *, collection_name: str = "default_space") -> None:
    """增量更新单文件（watchdog / 侧边栏上传触发），入库成功后更新空间专属状态文件。"""
    if path_should_skip(file_path):
        return

    doc = build_index_document(file_path, cfg)
    doc_id = str(file_path.resolve())

    try:
        index.delete_ref_doc(doc_id, delete_from_docstore=True)
    except Exception:
        pass

    if doc is None:
        return

    nodes = _nodes_from_built_document(doc, file_path, cfg)
    if not nodes:
        return

    ok = safe_insert_nodes(index, nodes, f"增量入库 ({file_path.name})")

    state = _load_index_state(collection_name)
    try:
        mtime = os.path.getmtime(str(file_path.resolve()))
    except OSError:
        mtime = 0.0
    _update_state_entry(state, doc_id, mtime, ok)
    _save_index_state(state, collection_name)


def initial_scan(index: VectorStoreIndex, cfg: AppConfig,
                 collection_name: str = "default_space",
                 progress_callback: Callable[[int, int], None] | None = None) -> int:
    """
    基于 mtime 的增量扫描，状态文件按空间隔离。
    progress_callback(current_batch, total_batches) 每批次完成后调用。
    返回本次入库的节点总数。
    """
    paths = [p for p in walk_workspace_files(cfg.workspaces)]
    if not paths:
        return 0

    file_mtimes: dict[str, float] = {}
    for p in paths:
        try:
            file_mtimes[str(p.resolve())] = os.path.getmtime(str(p))
        except OSError:
            continue

    current_keys = set(file_mtimes.keys())

    state = _load_index_state(collection_name)
    state = _prune_stale_entries(state, current_keys)

    pending: list[Path] = []
    skipped = 0
    for p in paths:
        key = str(p.resolve())
        cur_mtime = file_mtimes.get(key)
        if cur_mtime is None:
            continue
        entry = state.get(key)
        if isinstance(entry, dict) and entry.get("mtime") == cur_mtime and entry.get("in_index"):
            skipped += 1
            continue
        pending.append(p)

    if skipped:
        print(f"[第二大脑] mtime 缓存命中，跳过 {skipped} 个未变化文件")

    if not pending:
        _save_index_state(state, collection_name)
        return 0

    print(f"[第二大脑] 待扫描 {len(pending)} 个新/变更文件…")

    workers = _initial_scan_max_workers(len(pending))
    file_results: dict[str, list] = {}

    with ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {executor.submit(_ingest_worker_task, p, cfg): p for p in pending}
        for fut in as_completed(future_map):
            try:
                r = fut.result()
            except Exception:
                continue
            if r is not None:
                doc_id, nodes = r
                file_results[doc_id] = nodes

    if not file_results:
        _save_index_state(state, collection_name)
        return 0

    items = list(file_results.items())
    total_nodes = 0
    total_batches = (len(items) + INITIAL_SCAN_BATCH - 1) // INITIAL_SCAN_BATCH

    for batch_idx in range(0, len(items), INITIAL_SCAN_BATCH):
        batch = items[batch_idx:batch_idx + INITIAL_SCAN_BATCH]
        batch_num = batch_idx // INITIAL_SCAN_BATCH + 1
        batch_nodes: list = []

        for doc_id, nodes in batch:
            try:
                index.delete_ref_doc(doc_id, delete_from_docstore=True)
            except Exception:
                pass
            batch_nodes.extend(nodes)

        ok = safe_insert_nodes(index, batch_nodes, f"批次 {batch_num}/{total_batches}")
        if ok:
            total_nodes += len(batch_nodes)

        for doc_id, _ in batch:
            _update_state_entry(state, doc_id, file_mtimes[doc_id], ok)

        _save_index_state(state, collection_name)

        if progress_callback:
            progress_callback(batch_num, total_batches)

    if torch.cuda.is_available():
        torch.cuda.empty_cache()

    return total_nodes


# ============================================================================
# LLM 构造（对话模型，非 Embedding）
# ============================================================================

def get_llm_instance(selected_llm_id: str, cfg: AppConfig) -> tuple[LLM, str | None]:
    """
    按侧边栏选择的模型 ID 构造对话 LLM。
    返回 (llm, 若发生回退则为人可读提示文案，否则 None)。
    """
    sid = (selected_llm_id or "").strip() or cfg.llm_model

    if sid in ("qwen-plus", "qwen-max"):
        key = (cfg.dashscope_api_key or "").strip()
        if not key:
            msg = "未配置 DASHSCOPE_API_KEY，已回退到 Gemini。"
            print(f"[第二大脑] {msg}")
            llm = _make_gemini_llm(cfg)
            Settings.llm = llm
            return llm, msg
        try:
            os.environ["DASHSCOPE_API_KEY"] = key
            from llama_index.llms.dashscope import DashScope

            llm = DashScope(model_name=sid, api_key=key)
            Settings.llm = llm
            return llm, None
        except Exception as e:
            msg = f"千问加载失败（{e!s}），已回退到 Gemini。"
            print(f"[第二大脑] {msg}")
            llm = _make_gemini_llm(cfg)
            Settings.llm = llm
            return llm, msg

    model = sid if sid.startswith("gemini") else cfg.llm_model
    llm = _make_gemini_llm(cfg, model_name=model)
    Settings.llm = llm
    return llm, None


# ============================================================================
# 聊天引擎（BM25+Vector 混合检索 + 本地 GPU Rerank）
# ============================================================================

def format_sources_md(source_nodes: Iterable) -> str:
    lines = []
    seen: set[str] = set()
    for sn in source_nodes:
        meta = getattr(sn, "metadata", {}) or {}
        fp = meta.get("file_path") or meta.get("file_name")
        if fp and fp not in seen:
            seen.add(str(fp))
            name = meta.get("file_name") or Path(str(fp)).name
            ctype = meta.get("content_type", "")
            tag = "（图片索引）" if ctype == "image_caption" else ""
            lines.append(f"- **`{name}`**{tag}  ")
            lines.append(f"  `{fp}`  ")
            try:
                text = sn.get_text()
                if text:
                    snippet = text[:300].replace("`", "\\`")
                    ext = Path(name).suffix.lstrip(".") or "text"
                    lines.append(f"  ```{ext}")
                    lines.append(f"  {snippet}...")
                    lines.append(f"  ```")
            except Exception:
                pass
            lines.append("")
    if not lines:
        return "_（本次检索未命中带路径的节点）_"
    return "\n".join(lines)


def build_memory_for_completed_turns(cfg: AppConfig) -> ChatMemoryBuffer:
    return ChatMemoryBuffer.from_defaults(token_limit=cfg.memory_token_limit)


def build_chat_engine(
    index: VectorStoreIndex,
    llm: LLM,
    cfg: AppConfig,
    memory: ChatMemoryBuffer,
    collection_name: str = "default_space",
) -> BaseChatEngine:
    device = "cuda:0" if torch.cuda.is_available() else "cpu"
    reranker = SentenceTransformerRerank(
        model=cfg.local_rerank_model,
        top_n=cfg.top_k,
        device=device,
    )

    # BM25 + Vector 混合检索 (RRF 融合)
    vector_retriever = index.as_retriever(similarity_top_k=cfg.top_k)

    # BM25 缓存：按 index_state mtime 判断是否需要重建（避免每次对话 O(N) 词表构建）
    state_path = _index_state_path(collection_name)
    state_mtime = state_path.stat().st_mtime if state_path.exists() else 0.0

    cached = _BM25_CACHE.get(collection_name)
    if cached is not None and cached[0] == state_mtime:
        bm25_retriever = cached[1]
    else:
        try:
            bm25_retriever = BM25Retriever.from_defaults(
                index=index,
                similarity_top_k=cfg.top_k,
                tokenizer=code_aware_tokenize,
            )
            _BM25_CACHE[collection_name] = (state_mtime, bm25_retriever)
        except Exception:
            bm25_retriever = None

    if bm25_retriever is not None:
        retriever = QueryFusionRetriever(
            retrievers=[vector_retriever, bm25_retriever],
            similarity_top_k=cfg.top_k,
            num_queries=1,
            mode="reciprocal_rank_fusion",
        )
    else:
        retriever = vector_retriever

    return CondensePlusContextChatEngine.from_defaults(
        retriever=retriever,
        llm=llm,
        memory=memory,
        system_prompt=EXPERT_SYSTEM_PROMPT,
        node_postprocessors=[reranker],
        streaming=True,
    )


# ============================================================================
# 启动引导
# ============================================================================

def bootstrap_index(collection_name: str = "default_space") \
        -> tuple[AppConfig, HuggingFaceEmbedding, VectorStoreIndex]:
    """
    初始化本地 Embedding + 指定空间的 Chroma 向量库。
    返回 (cfg, embed_model, index)。
    """
    cfg = get_config()

    embed_model = _make_local_embed(cfg)
    Settings.embed_model = embed_model
    Settings.chunk_size = cfg.chunk_size
    Settings.chunk_overlap = cfg.chunk_overlap
    Settings.llm = _make_gemini_llm(cfg)

    index = _get_or_create_index(cfg, embed_model, collection_name)

    cleanup_ghost_nodes(index, collection_name)
    total = initial_scan(index, cfg, collection_name)
    print(f"[第二大脑] 空间 [{collection_name}] 启动扫描完成 — 入库 {total} nodes")

    return cfg, embed_model, index
